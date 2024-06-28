import os
import logging
import sqlite3
import asyncio
import requests
import shutil
import platform
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Query
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document as DocxDocument
from aiopytesseract import image_to_string
import fitz
from llama_index.llms.ollama import Ollama
from typing import Optional
from PIL import Image
import aiofiles

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Initialize FastAPI app
app = FastAPI(
    title="File Summarization API",
    description="An API for uploading, summarizing, and downloading files.",
    version="1.0.0",
)
app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"]
)

# Initialize LLaMA model with Ollama
llm = Ollama(model="llama3", request_timeout=120.0)

# Directory setup
os.makedirs(os.path.join("uploads"), exist_ok=True)
os.makedirs(os.path.join("output-files", "docx"), exist_ok=True)
os.makedirs(os.path.join("output-files", "raw-content"), exist_ok=True)

# Database setup
conn = sqlite3.connect(os.path.join("summarization.db"), check_same_thread=False)
c = conn.cursor()
c.execute(
    """CREATE TABLE IF NOT EXISTS files
             (id INTEGER PRIMARY KEY, original_filename TEXT, saved_path TEXT, summarized_filename TEXT, status TEXT, raw_content_path TEXT)"""
)
conn.commit()


class FileRecord(BaseModel):
    id: int
    original_filename: str
    saved_path: str
    summarized_filename: str
    status: str
    raw_content_path: str


class SummarizeRequest(BaseModel):
    id: int


# Initialize queue for task management
task_queue = asyncio.Queue()


# Function to find the Tesseract executable path
def find_tesseract_path():
    if platform.system() == "Darwin":  # Check if the system is MacOS
        tesseract_path = "/opt/homebrew/Cellar/tesseract/5.3.4_1/share/tessdata"
    else:
        tesseract_path = shutil.which("tesseract")
        if not tesseract_path:
            raise RuntimeError(
                "Tesseract executable not found. Please install Tesseract."
            )
    return tesseract_path


# Function to get the tessdata directory path
def get_tessdata_dir():
    tesseract_path = find_tesseract_path()
    tessdata_dir = os.path.join(os.path.dirname(tesseract_path), "tessdata")
    os.makedirs(tessdata_dir, exist_ok=True)
    return tessdata_dir


# Function to download tessdata files
def download_tessdata(lang_code):
    tessdata_dir = get_tessdata_dir()
    lang_file = f"{lang_code}.traineddata"
    lang_file_path = os.path.join(tessdata_dir, lang_file)

    if not os.path.exists(lang_file_path):
        url = (
            f"https://raw.githubusercontent.com/tesseract-ocr/tessdata/main/{lang_file}"
        )
        response = requests.get(url)
        if response.status_code == 200:
            with open(lang_file_path, "wb") as f:
                f.write(response.content)
            logging.info(f"Downloaded {lang_file} to {lang_file_path}")
        else:
            logging.error(f"Failed to download {lang_file} from {url}")
    else:
        logging.info(f"{lang_file} already exists at {lang_file_path}")


# Download necessary tessdata files
download_tessdata("fas")
download_tessdata("eng")


@app.post("/upload", summary="Upload a new file", response_model=dict)
async def upload_file(file: UploadFile = File(...)):
    """
    Upload a file for summarization.
    """
    original_filename = file.filename
    saved_path = os.path.join("uploads", original_filename)
    async with aiofiles.open(saved_path, "wb") as f:
        await f.write(await file.read())

    query = "INSERT INTO files (original_filename, saved_path, summarized_filename, status, raw_content_path) VALUES (?, ?, ?, ?, ?)"
    params = (original_filename.lower(), saved_path, "", "uploaded", "")
    c.execute(query, params)
    conn.commit()
    new_id = c.lastrowid
    return {"id": new_id}


@app.post(
    "/summarize", summary="Start summarization of an uploaded file", response_model=dict
)
async def summarize_file(request: SummarizeRequest, background_tasks: BackgroundTasks):
    """
    Start the summarization process for an uploaded file.
    """
    c.execute(
        "SELECT original_filename, saved_path, summarized_filename, status FROM files WHERE id = ?",
        (request.id,),
    )
    file_record = c.fetchone()
    if not file_record:
        raise HTTPException(status_code=404, detail="File not found")

    original_filename, saved_path, summarized_filename, status = file_record

    if status == "uploaded":
        c.execute("UPDATE files SET status = ? WHERE id = ?", ("queued", request.id))
        conn.commit()
        await task_queue.put((request.id, saved_path, original_filename))
        background_tasks.add_task(process_queue)
        return {"message": "Summarization queued"}
    else:
        raise HTTPException(
            status_code=400, detail=f"File status is {status}, not 'uploaded'"
        )


@app.get("/status", summary="Check the status of the summarization job")
async def status(
    task_id: int = Query(..., description="Task ID of the summarization job")
):
    """
    Get the current status of the summarization job.
    """

    async def generate(task_id: int):
        while True:
            c.execute("SELECT status FROM files WHERE id = ?", (task_id,))
            file_record = c.fetchone()
            if file_record:
                (status,) = file_record
                logging.info(f"Current status for task {task_id}: {status}")
                yield f"data: {status}\n\n"
                if status in ["completed", "error"]:
                    break
            await asyncio.sleep(1)

    return StreamingResponse(generate(task_id), media_type="text/event-stream")


@app.get(
    "/search",
    summary="Search for the summarized file using either the ID or the filename",
)
async def search_file(
    id: Optional[int] = Query(None, description="File ID to search for the summary"),
    filename: Optional[str] = Query(
        None, description="Filename (including extension) to search for the summary"
    ),
):
    """
    Search for the summarized file either by ID or by filename.
    """
    if bool(id) == bool(filename):
        raise HTTPException(
            status_code=400, detail="Provide either 'id' or 'filename', but not both"
        )

    # Log the request for debugging purposes
    logging.debug(f"Search invoked with id: {id}, filename: {filename}")

    if id:
        c.execute(
            "SELECT original_filename, saved_path, summarized_filename, status, raw_content_path FROM files WHERE id = ?",
            (id,),
        )
    else:
        c.execute(
            "SELECT original_filename, saved_path, summarized_filename, status, raw_content_path FROM files WHERE original_filename = ?",
            (filename.lower(),),
        )

    file_record = c.fetchone()

    if not file_record:
        raise HTTPException(status_code=404, detail="File not found")

    original_filename, saved_path, summarized_filename, status, raw_content_path = (
        file_record
    )
    if status != "completed":
        raise HTTPException(
            status_code=400, detail=f"File summarization status: {status}"
        )

    summary_path = os.path.join("output-files", "docx", summarized_filename)
    raw_text_path = raw_content_path

    # Log path to the summary and raw text for debugging purposes
    logging.debug(f"Summary path resolved to: {summary_path}")
    logging.debug(f"Raw text path resolved to: {raw_text_path}")

    if os.path.exists(summary_path):
        try:
            return {
                "original_filename": original_filename,
                "summarized_filename": summarized_filename,
                "raw_content_path": raw_content_path,
            }
        except Exception as e:
            logging.exception("Error reading summarized file")
            raise HTTPException(status_code=500, detail=str(e))
    else:
        raise HTTPException(status_code=404, detail="Summarized file not found")


@app.post(
    "/update_db_schema",
    summary="Update the database schema to ensure it matches current requirements",
    response_model=dict,
)
async def update_db_schema():
    """
    Update the database schema if necessary.
    """
    try:
        c.execute("ALTER TABLE files ADD COLUMN raw_content_path TEXT")
        conn.commit()
    except sqlite3.OperationalError:
        pass  # Column already exists

    return {"message": "Database schema updated successfully"}


@app.get("/download", summary="Download the summarized file")
async def download_file(
    id: Optional[int] = Query(None, description="File ID to download the summary"),
    filename: Optional[str] = Query(
        None, description="Filename (including extension) to download the summary"
    ),
):
    """
    Download the summarized file either by ID or by filename.
    """
    if bool(id) == bool(filename):
        raise HTTPException(
            status_code=400, detail="Provide either 'id' or 'filename', but not both"
        )

    if id:
        c.execute(
            "SELECT original_filename, summarized_filename, status FROM files WHERE id = ?",
            (id,),
        )
    else:
        c.execute(
            "SELECT original_filename, summarized_filename, status FROM files WHERE original_filename = ?",
            (filename.lower(),),
        )

    file_record = c.fetchone()

    if not file_record:
        raise HTTPException(status_code=404, detail="File not found")

    original_filename, summarized_filename, status = file_record
    if status != "completed":
        raise HTTPException(
            status_code=400, detail=f"File summarization status: {status}"
        )

    summary_path = os.path.join("output-files", "docx", summarized_filename)
    if os.path.exists(summary_path):
        return FileResponse(
            summary_path,
            filename=summarized_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        raise HTTPException(status_code=404, detail="Summarized file not found")


@app.get("/download_raw", summary="Download the raw content file")
async def download_raw_file(
    id: Optional[int] = Query(None, description="File ID to download the raw content"),
    filename: Optional[str] = Query(
        None, description="Filename (including extension) to download the raw content"
    ),
):
    """
    Download the raw content file either by ID or by filename.
    """
    if bool(id) == bool(filename):
        raise HTTPException(
            status_code=400, detail="Provide either 'id' یا 'filename', but not both"
        )

    if id:
        c.execute(
            "SELECT original_filename, raw_content_path, status FROM files WHERE id = ?",
            (id,),
        )
    else:
        c.execute(
            "SELECT original_filename, raw_content_path, status FROM files WHERE original_filename = ?",
            (filename.lower(),),
        )

    file_record = c.fetchone()

    if not file_record:
        raise HTTPException(status_code=404, detail="File not found")

    original_filename, raw_content_path, status = file_record
    if status != "completed":
        raise HTTPException(
            status_code=400, detail=f"File summarization status: {status}"
        )

    if os.path.exists(raw_content_path):
        return FileResponse(
            raw_content_path,
            filename=os.path.basename(raw_content_path),
            media_type="text/plain",
        )
    else:
        raise HTTPException(status_code=404, detail="Raw content file not found")


async def process_queue():
    while not task_queue.empty():
        file_id, saved_path, original_filename = await task_queue.get()
        logging.info(f"Processing file with ID: {file_id}")
        c.execute("UPDATE files SET status = ? WHERE id = ?", ("processing", file_id))
        conn.commit()
        await process_summarization(file_id, saved_path, original_filename)
        task_queue.task_done()


async def process_summarization(file_id, saved_path, original_filename):
    extension = saved_path.split(".")[-1].lower()
    summarized_filename = f"{os.path.splitext(original_filename)[0]}_summary.docx"
    raw_content_filename = f"{os.path.splitext(original_filename)[0]}_raw.txt"

    text = ""
    if extension == "pdf":
        c.execute(
            "UPDATE files SET status = ? WHERE id = ?", ("extracting_text", file_id)
        )
        conn.commit()
        text = await extract_text_from_pdf(saved_path)
    elif extension == "docx":
        c.execute(
            "UPDATE files SET status = ? WHERE id = ?", ("extracting_text", file_id)
        )
        conn.commit()
        text = extract_text_from_docx(saved_path)
    elif extension == "txt":
        c.execute(
            "UPDATE files SET status = ? WHERE id = ?", ("extracting_text", file_id)
        )
        conn.commit()
        text = extract_text_from_txt(saved_path)
    elif extension in ["jpg", "jpeg", "png"]:
        c.execute(
            "UPDATE files SET status = ? WHERE id = ?", ("extracting_text", file_id)
        )
        conn.commit()
        text = await extract_text_from_image(saved_path)
    else:
        c.execute("UPDATE files SET status = ? WHERE id = ?", ("error", file_id))
        conn.commit()
        return

    logging.debug(f"Extracted text: {text}")

    if text:
        c.execute("UPDATE files SET status = ? WHERE id = ?", ("summarizing", file_id))
        conn.commit()
        chunk_size = 5000
        text_chunks = chunk_text(text, chunk_size)
        chunk_summaries = await asyncio.gather(
            *[summarize_text(llm, chunk) for chunk in text_chunks]
        )

        final_summary = "\n\n".join(chunk_summaries)

        output_dir = os.path.join("output-files", "docx")
        os.makedirs(output_dir, exist_ok=True)
        output_file_path = os.path.join(output_dir, summarized_filename)

        raw_content_dir = os.path.join("output-files", "raw-content")
        os.makedirs(raw_content_dir, exist_ok=True)
        raw_content_file_path = os.path.join(raw_content_dir, raw_content_filename)

        logging.info(f"Saving summary to: {output_file_path}")
        logging.info(f"Saving raw content to: {raw_content_file_path}")

        save_summary(final_summary, output_file_path)
        save_raw_content(text, raw_content_file_path)

        c.execute(
            "UPDATE files SET status = ?, summarized_filename = ?, raw_content_path = ? WHERE id = ?",
            ("completed", summarized_filename, raw_content_file_path, file_id),
        )
        conn.commit()
    else:
        c.execute("UPDATE files SET status = ? WHERE id = ?", ("error", file_id))
        conn.commit()


async def extract_text_from_pdf(pdf_path):
    try:
        text = ""
        document = fitz.open(pdf_path)
        for page_num in range(len(document)):
            page = document[page_num]
            # Convert PDF page to image for OCR extraction
            pix = page.get_pixmap()
            image_path = f"temp_image_{page_num+1}.png"
            pix.save(image_path)
            ocr_text = await image_to_string(image_path, lang="fas+eng")
            logging.debug(f"OCR text from page {page_num + 1}: {ocr_text}")
            text += ocr_text + "\n\n"
            os.remove(image_path)  # Remove the temporary image file
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {str(e)}")
        return ""


def extract_text_from_docx(docx_path):
    try:
        doc = DocxDocument(docx_path)
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {str(e)}")
        return ""


def extract_text_from_txt(txt_path):
    try:
        with open(txt_path, "r", encoding="utf-8") as file:
            text = file.read()
        return text
    except Exception as e:
        logging.error(f"Error extracting text from TXT: {str(e)}")
        return ""


async def extract_text_from_image(image_path):
    try:
        image = Image.open(image_path)
        ocr_text = await image_to_string(image, lang="fas+eng")
        logging.debug(f"OCR text from image {image_path}: {ocr_text}")
        return ocr_text
    except Exception as e:
        logging.error(f"Error extracting text from image: {str(e)}")
        return ""


def chunk_text(text, chunk_size):
    return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]


async def summarize_text(llm, text, max_retries=3):
    attempt = 0
    while attempt < max_retries:
        try:
            summarize_prompt = (
                "Please provide a concise summary of the following text, just send summary, don't say 'Here is a concise summary:' or other related phrases too.:\n\n"
                f"{text}\n\n"
                "Summary:"
            )
            response = llm.complete(prompt=summarize_prompt)
            summary_content = str(response)
            return summary_content.strip()
        except Exception as e:
            logging.error(
                f"Error summarizing text (attempt {attempt + 1}/{max_retries}): {str(e)}"
            )
            if "503" in str(e):
                attempt += 1
                await asyncio.sleep(2**attempt)
            else:
                break
    return ""


def save_summary(summary, output_file):
    try:
        doc = DocxDocument()
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)
        doc.save(output_file)
        logging.info(f"Summary saved to {output_file}")
    except Exception as e:
        logging.error(f"Error saving summary: {str(e)}")


def save_raw_content(raw_content, output_file):
    try:
        with open(output_file, "w", encoding="utf-8") as file:
            file.write(raw_content)
        logging.info(f"Raw content saved to {output_file}")
    except Exception as e:
        logging.error(f"Error saving raw content: {str(e)}")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=3000)
